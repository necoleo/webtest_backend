from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html
import pandas as pd

from demand_vector.src.config.loader import loadConfig
from demand_vector.src.embedding.zhipu_model import ZhiPu4
from demand_vector.src.indexing.vector_indexing import VectorIndexing
from demand_vector.src.search.sentence_search import SimilarSentenceSearch


# 加载配置
config = loadConfig().load_config()

# 执行任务
run_task = config['EXECUTE_TASK']

# 加载总结提炼设定长度值
SUMMARY_LENGTH = config['SUMMARY_LENGTH']

if run_task == '1':
    # 执行拆分需求
    # 执行拆分需求转为向量
    # 将向量写入向量数据库
    # 将向量数据与文本数据通过MySQL关联
    vector_indexing = VectorIndexing()
    vector_indexing.process_texts()

if run_task == '2':
    # 获取所有文档的热词
    # 从数据库中查询所有的文本
    faiss_index_path = "../data/processed/" + config['FAISS_DB_PATH']
    # 相似句子查询实例
    similar_sentence_search = SimilarSentenceSearch(faiss_index_path)
    # 创建智谱实例
    zhipu_4 = ZhiPu4()

    information_document = similar_sentence_search.get_information_document()
    information_document_list = zhipu_4.split_list_by_text_length(information_document, SUMMARY_LENGTH)
    response = zhipu_4.concurrent_zhipuai_request_text(information_document_list)

    # 相似度阈值
    similarity_threshold = config['SIMILARITY_THRESHOLD']
    vdb = similar_sentence_search.vectordb

    # 创建Word文档
    doc_filename = '../data/processed/requirements.docx'
    doc = Document()
    doc.add_heading('Requirements', 0)  # 添加标题

    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)
    font.name = 'Arial'

    elements = []
    for data in response:
        for text in eval(data):
            # 执行相似度查询
            similar_texts = similar_sentence_search.search(text, similarity_threshold)
            print(similar_texts)
            # 将数据存入数据库
            vdb.create_needs_text(text, str(similar_texts))

            # 将HTML实体转换回原始字符
            text = html.unescape(text)
            similar_texts = html.unescape(str(similar_texts))

            # 添加文本到Word文档
            doc.add_paragraph(text, style=style)
            doc.add_paragraph(similar_texts, style=style)

            # 添加一些空白
            doc.add_paragraph('', style=style)

    # 保存Word文档
    doc.save(doc_filename)

if run_task == '3':

    # 读取Excel文件
    df = pd.read_excel('C:\\Users\\walker\\Desktop\\collected_needs_data.xlsx')

    # 使用groupby按照指定的列进行分组，并将每组内的case, operation_steps, expected_results合并成字典列表
    grouped = df.groupby(['needs_info', 'module_name', 'key', 'need']).apply(lambda x: [
        {
            'case': case,
            'operation_steps': operation_steps,
            'expected_results': expected_results
        }
        for case, operation_steps, expected_results in zip(x['case'], x['operation_steps'], x['expected_results'])
    ]).reset_index(name='data')

    # 转换为最终所需的字典格式
    result_dict = {
        (row['needs_info'], row['module_name'], row['key'], row['need']): row['data']
        for _, row in grouped.iterrows()
    }

    jsonl_data = []
    # 打印结果
    for key, value in result_dict.items():
        processed_value = [
            {
                'testpoint': item['case'],
                'operation': item['operation_steps'].replace('  ', ''),
                'expectedresult': item['expected_results'].replace('  ', '')
            }
            for item in value
        ]

        try:
            needs_info = key[0].replace('\\', '')
        except:
            needs_info = key[0]

        jsonl_data.append(
            (f"测试模块为：{key[1]}，测试方向为：{key[2]}，指定需要生成测试用例的需求为：{key[3]}", processed_value))

    # print(jsonl_data)
    vector_indexing = VectorIndexing()
    vector_indexing.case_embedding(jsonl_data)