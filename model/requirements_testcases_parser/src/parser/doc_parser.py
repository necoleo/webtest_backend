import os
import re
import zipfile

import jieba
from docx import Document


class WordDocumentParser:

    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.document = Document(doc_path)
        self.img_temp_dir = "../../data/processed/img_temp"

    def extract_all_images(self):
        # 确保输出目录存在
        if not os.path.exists(self.img_temp_dir):
            os.makedirs(self.img_temp_dir, exist_ok=True)

        extracted_images = []

        with zipfile.ZipFile(self.doc_path) as docx_zip:
            # 查找所有媒体文件
            media_files = [f for f in docx_zip.namelist()
                           if f.startswith('word/media/')]

            # 提取每个媒体文件
            for media_file in media_files:
                # 获取文件名
                filename = os.path.basename(media_file)
                # 构建输出路径
                output_path = os.path.join(self.img_temp_dir, filename)
                # 提取文件
                with docx_zip.open(media_file) as source, open(output_path, 'wb') as target:
                    target.write(source.read())
                extracted_images.append(output_path)

        return extracted_images

    def split_sentences(self, text):
        """
        根据正则匹配划分句子
        :param text:
        :return:
        """
        text_data = re.split(r'(?<=[。！？!?])', text)
        sentences_list = []
        for data in text_data:
            if data == '':
                continue
            if len(data) < 4:
                continue
            sentences_list.append(data)
        return sentences_list

    def get_document_paragraphs(self):
        """
        获取文档所有段落文本
        :return:
        """
        document_content = []
        file_name = os.path.basename(self.doc_path).split('.')[0]
        if len(self.document.paragraphs) < 1:
            return False, []
        else:
            for para in self.document.paragraphs:
                if para.text != '':
                    sentence_list = self.split_sentences(para.text)
                    for sentence in sentence_list:
                        if re.search("需求背景", sentence) or re.search("需求说明", sentence):
                            if len(sentence) < 8:
                                continue
                            else:
                                sentence = file_name + "_" +sentence
                                document_content.append(sentence)
                        else:
                            sentence = file_name + "_" +sentence
                            document_content.append(sentence)
            return True, document_content

if __name__ == '__main__':
    doc = WordDocumentParser('C:/Users/92700/Desktop/tencent_meeting.docx')
    str = "【需求背景】"
    print(len(str))
    doc.extract_all_images()